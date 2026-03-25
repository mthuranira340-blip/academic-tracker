from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


OUTPUT_FILE = "University_Academic_Tracker_Documentation.docx"


def add_heading(document, text, level=1):
    document.add_heading(text, level=level)


def add_bullets(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document, items):
    for item in items:
        document.add_paragraph(item, style="List Number")


def build_document():
    document = Document()

    title = document.add_heading("University Academic Tracker Documentation", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = document.add_paragraph(
        "Complete project documentation for the Flask, MySQL, Bootstrap, and Chart.js application."
    )
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    add_heading(document, "1. Project Overview")
    document.add_paragraph(
        "University Academic Tracker is a full-stack academic management system built for students and parents. "
        "It helps manage units, grades, GPA, progress charts, payment prompts, AI-assisted study support, "
        "and calendar-based reminders for academic and co-curricular activities."
    )

    add_heading(document, "2. Main Objectives")
    add_bullets(
        document,
        [
            "Provide secure student and parent authentication.",
            "Track units, grades, and GPA in a structured MySQL database.",
            "Visualize academic progress using charts.",
            "Offer parent read-only access to student performance.",
            "Display academic-calendar payment prompts and pricing.",
            "Support assignments and course notes through an AI assistant page.",
            "Manage academic and co-curricular activities with 24-hour reminders.",
        ],
    )

    add_heading(document, "3. Technology Stack")
    add_bullets(
        document,
        [
            "Backend: Python Flask",
            "Database: MySQL",
            "ORM: Flask-SQLAlchemy",
            "Authentication: Flask-Login and Werkzeug password hashing",
            "Forms and validation: Flask-WTF / WTForms",
            "Frontend: HTML, CSS, JavaScript, Bootstrap 5",
            "Charts: Chart.js",
            "PDF Export: ReportLab",
            "Word Documentation Output: python-docx",
        ],
    )

    add_heading(document, "4. Core Features")
    add_bullets(
        document,
        [
            "User registration and login for students and parents.",
            "Role-based access control with parent read-only restrictions.",
            "CRUD operations for units.",
            "CRUD operations for grades.",
            "Automatic GPA calculation.",
            "Academic performance visualization over time.",
            "Low-grade notifications.",
            "PDF export for academic reports.",
            "Dark mode toggle.",
            "Academic-calendar payment prompt with service charges in USD.",
            "PayPal payment destination display.",
            "AI assistant for assignments and unit-specific notes.",
            "Calendar planner with academic and co-curricular activities.",
            "24-hour reminders for upcoming activities.",
        ],
    )

    add_heading(document, "5. User Roles")
    document.add_paragraph("Student Role")
    add_bullets(
        document,
        [
            "Register and log in.",
            "Add, edit, and delete units.",
            "Add, edit, and delete grades.",
            "View GPA, charts, alerts, and payment prompts.",
            "Use the AI assistant for assignments and notes.",
            "Add and remove calendar activities.",
        ],
    )
    document.add_paragraph("Parent Role")
    add_bullets(
        document,
        [
            "Register and link to a student account.",
            "Log in and view student grades, GPA, charts, and reminders.",
            "Access the AI assistant page in read-only study-support mode.",
            "Cannot create, edit, or delete academic data.",
        ],
    )

    add_heading(document, "6. Database Design")
    document.add_paragraph("The application uses the following MySQL tables:")
    add_bullets(
        document,
        [
            "users: stores username, email, hashed password, and role.",
            "units: stores unit name, unit code, course, and level of study.",
            "grades: stores student grades linked to users and units.",
            "parent_student_links: connects parent accounts to student accounts.",
            "activities: stores academic and co-curricular events with 24-hour reminders.",
        ],
    )

    add_heading(document, "7. GPA Logic")
    document.add_paragraph("Grade points are mapped as follows:")
    add_bullets(
        document,
        [
            "A = 4.0",
            "B+ = 3.5",
            "B = 3.0",
            "C+ = 2.5",
            "C = 2.0",
            "D = 1.0",
            "E = 0.0",
        ],
    )
    document.add_paragraph(
        "GPA is calculated by summing all grade points and dividing by the number of grade entries."
    )

    add_heading(document, "8. Payment System")
    document.add_paragraph(
        "The dashboard includes an academic-calendar-based payment prompt. Each semester includes service charges "
        "in USD and a total amount due. The payment destination is displayed as a PayPal account."
    )
    add_bullets(
        document,
        [
            "Payment account: mthuranira340@gmail.com",
            "Service pricing is shown per semester.",
            "The AI assistant pricing model is set at USD 30 for every 100 support hours.",
            "Dollar icons are used in the user interface to make payment sections clearer.",
        ],
    )

    add_heading(document, "9. AI Assistant")
    document.add_paragraph(
        "The AI assistant page helps students with assignments and course revision. Users can select a unit, "
        "choose an assignment type, enter the number of support hours, and submit a study request."
    )
    add_bullets(
        document,
        [
            "Unit filtering for notes.",
            "Best notes per selected unit.",
            "Assignment guidance output.",
            "Pricing calculator for study-support time.",
            "Built-in offline helper when a live API key is not configured.",
            "Optional live AI support through an OpenAI API key.",
        ],
    )

    add_heading(document, "10. Calendar and Reminders")
    document.add_paragraph(
        "Students can add academic calendar items and co-curricular activities from the dashboard. "
        "Each activity stores the title, category, date/time, description, and a reminder time set exactly 24 hours earlier."
    )
    add_bullets(
        document,
        [
            "Academic and co-curricular event categories.",
            "Date and time entry using a 24-hour format field.",
            "Reminder center for activities occurring within the next 24 hours.",
            "Timeline view for upcoming and completed events.",
            "Parents can view reminders but cannot edit activities.",
        ],
    )

    add_heading(document, "11. Project Structure")
    add_bullets(
        document,
        [
            "run.py: application entry point.",
            "config.py: configuration and environment settings.",
            "app/__init__.py: Flask app factory and setup.",
            "app/models.py: SQLAlchemy models.",
            "app/forms.py: WTForms form classes.",
            "app/services.py: GPA, payment, AI, notes, and reminder logic.",
            "app/routes/auth.py: registration, login, and logout routes.",
            "app/routes/main.py: dashboard, units, grades, assistant, PDF, and activities routes.",
            "app/templates/: HTML templates.",
            "app/static/css/style.css: custom styles.",
            "app/static/js/app.js: frontend interactivity and charts.",
            "schema.sql: MySQL schema definition.",
        ],
    )

    add_heading(document, "12. Setup Instructions")
    add_numbered(
        document,
        [
            "Install Python 3.12 or later.",
            "Create a virtual environment and activate it.",
            "Install dependencies with pip install -r requirements.txt.",
            "Create the MySQL database using schema.sql.",
            "Set the SECRET_KEY and DATABASE_URL environment variables.",
            "Run the app with python run.py.",
            "Open the application in a browser at http://127.0.0.1:5000.",
        ],
    )

    add_heading(document, "13. Environment Variables")
    add_bullets(
        document,
        [
            "SECRET_KEY: Flask secret key.",
            "DATABASE_URL: SQLAlchemy connection string for MySQL.",
            "LOW_GRADE_THRESHOLD: optional threshold for alerts.",
            "OPENAI_API_KEY: optional key for live AI assistant responses.",
            "OPENAI_MODEL: optional model selection for the assistant feature.",
        ],
    )

    add_heading(document, "14. Security Notes")
    add_bullets(
        document,
        [
            "Passwords are securely hashed using Werkzeug.",
            "Forms use CSRF protection through Flask-WTF.",
            "Role-based access blocks parents from modifying student data.",
            "Database relationships use foreign keys and cascade rules.",
        ],
    )

    add_heading(document, "15. Testing and Validation")
    add_bullets(
        document,
        [
            "Python modules compile successfully.",
            "Application routes have been exercised through the local development server.",
            "MySQL tables have been created and verified.",
            "The Flask development server is running locally during implementation.",
        ],
    )

    add_heading(document, "16. Suggested Future Improvements")
    add_bullets(
        document,
        [
            "Add email or SMS notifications for reminders and low grades.",
            "Allow multiple students per parent account.",
            "Support file uploads for class notes and assignment resources.",
            "Add admin analytics and institution-level reporting.",
            "Integrate direct online payment processing workflows.",
            "Add recurring calendar events and timetable import.",
        ],
    )

    add_heading(document, "17. Conclusion")
    document.add_paragraph(
        "University Academic Tracker is a beginner-friendly but feature-rich academic management application. "
        "It combines academic monitoring, payment prompts, AI-based study support, and event planning into one system."
    )

    document.save(OUTPUT_FILE)


if __name__ == "__main__":
    build_document()
